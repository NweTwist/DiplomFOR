#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Собирает ВКР из трёх глав в единый DOCX по ГОСТ Р 7.0.11-2011.

Правила форматирования:
  - Шрифт: Times New Roman 14 везде (текст, заголовки, таблицы, подписи)
  - Поля: левое 3 см, правое 1.5 см, верхнее 2 см, нижнее 2 см
  - Межстрочный интервал: 1.5
  - Абзацный отступ: 1.25 см, выравнивание по ширине
  - Нумерация страниц: внизу по центру
  - Заголовки глав: по центру, КАПСЛОК
  - Заголовки разделов: по левому краю, жирный
  - Название таблицы — ПОСЛЕ таблицы, по центру
  - Подпись рисунка — под рисунком, по центру
  - Введение — отдельный раздел перед главами
  - Содержание — автоматическое, обновляется при открытии в Word
"""

from __future__ import annotations
import re
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Константы
# ---------------------------------------------------------------------------
FONT      = "Times New Roman"
SZ_BODY   = Pt(14)
SZ_H1     = Pt(16)   # Глава
SZ_H2     = Pt(14)   # Раздел
SZ_H3     = Pt(14)   # Подраздел
SZ_H4     = Pt(14)   # Подподраздел
SZ_TABLE  = Pt(12)
SZ_CAP    = Pt(12)   # Подписи рисунков и таблиц

L_MARGIN  = Cm(3.0)
R_MARGIN  = Cm(1.5)
T_MARGIN  = Cm(2.0)
B_MARGIN  = Cm(2.0)
INDENT    = Cm(1.25)

CHAPTERS = [
    Path("Первая глава.md"),
    Path("Вторая глава.md"),
    Path("Третья глава.md"),
]
OUTPUT = Path("VKR_Anomaly_Detection.docx")

REFERENCES_TEXT = """\
[1] Chandola V., Banerjee A., Kumar V. Anomaly Detection: A Survey // ACM Computing Surveys. — 2009. — Vol. 41, № 3. — Статья 15.

[2] Sommer R., Paxson V. Outside the Closed World: On Using Machine Learning for Network Intrusion Detection // IEEE Symposium on Security and Privacy. — 2010.

[3] Zhao Y., Nasrullah Z., Li Z. PyOD: A Python Toolbox for Scalable Outlier Detection // Journal of Machine Learning Research. — 2019. — Vol. 20, № 96. — С. 1–7.

[4] Paxson V. Bro: A System for Detecting Network Intruders in Real-Time // Computer Networks. — 1999. — Vol. 31, № 23–24. — С. 2435–2463.

[5] Deri L., Martinelli M., Bujlow T., Cardigliano A. nDPI: Open-Source High-Speed Deep Packet Inspection // International Wireless Communications and Mobile Computing Conference (IWCMC). — 2014.

[6] Liu F. T., Ting K. M., Zhou Z.-H. Isolation Forest // IEEE International Conference on Data Mining (ICDM). — 2008.

[7] Mirsky Y. et al. Kitsune: An Ensemble of Autoencoders for Online Network Intrusion Detection // NDSS Symposium. — 2018.

[8] ГОСТ Р 7.0.11-2011. Диссертация и автореферат диссертации. Структура и правила оформления. — М.: Стандартинформ, 2012.
"""


# ---------------------------------------------------------------------------
# Низкоуровневые хелперы
# ---------------------------------------------------------------------------

def _rpr_font(run, size: Pt, bold=False, italic=False, color=None) -> None:
    """Выставляет шрифт Times New Roman 14 (или другой size) для run."""
    run.font.name  = FONT
    run.font.size  = size
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    try:
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"),    FONT)
        rFonts.set(qn("w:hAnsi"),    FONT)
        rFonts.set(qn("w:eastAsia"), FONT)
        rFonts.set(qn("w:cs"),       FONT)
        existing = rPr.find(qn("w:rFonts"))
        if existing is not None:
            rPr.remove(existing)
        rPr.insert(0, rFonts)
    except Exception:
        pass


def _pfmt(para, *,
          align=WD_ALIGN_PARAGRAPH.JUSTIFY,
          first_indent: Cm | None = INDENT,
          left_indent: Cm | None = None,
          space_before: Pt = Pt(0),
          space_after: Pt = Pt(0),
          line_spacing=WD_LINE_SPACING.ONE_POINT_FIVE) -> None:
    pf = para.paragraph_format
    pf.alignment         = align
    pf.space_before      = space_before
    pf.space_after       = space_after
    pf.line_spacing_rule = line_spacing
    pf.first_line_indent = first_indent
    if left_indent is not None:
        pf.left_indent = left_indent


def set_page_margins(doc: Document) -> None:
    for s in doc.sections:
        s.left_margin   = L_MARGIN
        s.right_margin  = R_MARGIN
        s.top_margin    = T_MARGIN
        s.bottom_margin = B_MARGIN


def add_page_numbers(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        if footer.paragraphs:
            para = footer.paragraphs[0]
        else:
            para = footer.add_paragraph()
        para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for tag, text in [("begin", None), ("instrText", "PAGE"), ("end", None)]:
            run = para.add_run()
            if tag == "instrText":
                instr = OxmlElement("w:instrText")
                instr.text = text
                run._r.append(instr)
            else:
                fld = OxmlElement("w:fldChar")
                fld.set(qn("w:fldCharType"), tag)
                run._r.append(fld)
            _rpr_font(run, SZ_BODY)


# ---------------------------------------------------------------------------
# Добавление элементов документа
# ---------------------------------------------------------------------------

def add_body(doc: Document, text: str, *,
             indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             bold=False, italic=False) -> None:
    para = doc.add_paragraph()
    _pfmt(para, align=align, first_indent=INDENT if indent else None)
    run = para.add_run(text)
    _rpr_font(run, SZ_BODY, bold=bold, italic=italic)


def add_heading_chapter(doc: Document, text: str) -> None:
    """Заголовок главы: Times New Roman 16, жирный, по центру, КАПСЛОК."""
    para = doc.add_paragraph()
    _pfmt(para, align=WD_ALIGN_PARAGRAPH.CENTER,
          first_indent=None, space_before=Pt(18), space_after=Pt(12))
    run = para.add_run(text.upper())
    _rpr_font(run, SZ_H1, bold=True)


def add_heading_section(doc: Document, text: str) -> None:
    """Заголовок раздела: Times New Roman 14, жирный, по левому краю."""
    para = doc.add_paragraph()
    _pfmt(para, align=WD_ALIGN_PARAGRAPH.LEFT,
          first_indent=None, space_before=Pt(14), space_after=Pt(6))
    run = para.add_run(text)
    _rpr_font(run, SZ_H2, bold=True)


def add_heading_subsection(doc: Document, text: str) -> None:
    """Подраздел: Times New Roman 14, жирный курсив, с абзацным отступом."""
    para = doc.add_paragraph()
    _pfmt(para, align=WD_ALIGN_PARAGRAPH.LEFT,
          first_indent=INDENT, space_before=Pt(10), space_after=Pt(4))
    run = para.add_run(text)
    _rpr_font(run, SZ_H3, bold=True, italic=True)


def add_heading_sub4(doc: Document, text: str) -> None:
    """Подподраздел: курсив, с абзацным отступом."""
    para = doc.add_paragraph()
    _pfmt(para, align=WD_ALIGN_PARAGRAPH.LEFT,
          first_indent=INDENT, space_before=Pt(8), space_after=Pt(2))
    run = para.add_run(text)
    _rpr_font(run, SZ_H4, italic=True)


def add_bullet_item(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    _pfmt(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
          first_indent=Cm(-0.5), left_indent=Cm(1.75))
    run = para.add_run(f"— {text}")
    _rpr_font(run, SZ_BODY)


def add_code(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    _pfmt(para, align=WD_ALIGN_PARAGRAPH.LEFT,
          first_indent=None, left_indent=INDENT,
          space_before=Pt(4), space_after=Pt(4),
          line_spacing=WD_LINE_SPACING.SINGLE)
    run = para.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(10)


def add_image_with_caption(doc: Document, img_path: Path, caption: str) -> None:
    """Рисунок + подпись под ним по центру."""
    # Рисунок
    img_para = doc.add_paragraph()
    _pfmt(img_para, align=WD_ALIGN_PARAGRAPH.CENTER,
          first_indent=None, space_before=Pt(8), space_after=Pt(2))
    if img_path.exists():
        try:
            img_para.add_run().add_picture(str(img_path), width=Cm(15.0))
        except Exception as e:
            run = img_para.add_run(f"[Ошибка вставки: {img_path.name}]")
            _rpr_font(run, SZ_BODY, italic=True)
    else:
        run = img_para.add_run(f"[Файл не найден: {img_path}]")
        _rpr_font(run, SZ_BODY, italic=True)
    # Подпись
    cap_para = doc.add_paragraph()
    _pfmt(cap_para, align=WD_ALIGN_PARAGRAPH.CENTER,
          first_indent=None, space_before=Pt(0), space_after=Pt(10))
    run = cap_para.add_run(caption)
    _rpr_font(run, SZ_CAP)


def _parse_md_table(lines: list[str]) -> tuple[list[str], list[list[str]]]:
    rows = []
    for line in lines:
        line = line.strip()
        if not line:
            continue
        # Разделитель |---|--- пропускаем
        inner = line.strip("|")
        if set(inner.replace(" ", "").replace(":", "")) <= set("-"):
            continue
        cells = [c.strip() for c in line.split("|") if c.strip() != ""]
        if cells:
            rows.append(cells)
    if not rows:
        return [], []
    return rows[0], rows[1:]


def add_table_with_caption(doc: Document, header: list[str],
                            data: list[list[str]], caption: str) -> None:
    """По ГОСТ: сначала таблица, затем название по центру."""
    n_cols = len(header)
    if n_cols == 0:
        return
    table = doc.add_table(rows=1 + len(data), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Заголовочная строка
    for i, cell_text in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _pfmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=None,
              space_before=Pt(2), space_after=Pt(2))
        run = p.add_run(cell_text)
        _rpr_font(run, SZ_TABLE, bold=True)
        # Серый фон заголовка
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "D9D9D9")
        cell._tc.get_or_add_tcPr().append(shd)

    # Строки данных
    for r_i, row_data in enumerate(data):
        for c_i, cell_text in enumerate(row_data):
            if c_i >= n_cols:
                break
            cell = table.rows[r_i + 1].cells[c_i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            _pfmt(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=None,
                  space_before=Pt(1), space_after=Pt(1))
            run = p.add_run(cell_text)
            _rpr_font(run, SZ_TABLE)

    # Название таблицы ПОСЛЕ таблицы, по центру
    cap_para = doc.add_paragraph()
    _pfmt(cap_para, align=WD_ALIGN_PARAGRAPH.CENTER,
          first_indent=None, space_before=Pt(4), space_after=Pt(10))
    run = cap_para.add_run(caption)
    _rpr_font(run, SZ_CAP, italic=True)


def add_inline_para(doc: Document, text: str, *,
                    indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    """Абзац с inline-разметкой **bold** и `code`."""
    para = doc.add_paragraph()
    _pfmt(para, align=align, first_indent=INDENT if indent else None)
    # Разбиваем на сегменты
    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`)')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = para.add_run(text[pos:m.start()])
            _rpr_font(run, SZ_BODY)
        token = m.group()
        if token.startswith("**"):
            run = para.add_run(token[2:-2])
            _rpr_font(run, SZ_BODY, bold=True)
        else:
            run = para.add_run(token[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(12)
        pos = m.end()
    if pos < len(text):
        run = para.add_run(text[pos:])
        _rpr_font(run, SZ_BODY)


# ---------------------------------------------------------------------------
# Содержание (TOC) через поле Word
# ---------------------------------------------------------------------------

def add_toc(doc: Document) -> None:
    """Вставляет автоматическое содержание через поле TOC."""
    add_heading_chapter(doc, "СОДЕРЖАНИЕ")

    para = doc.add_paragraph()
    _pfmt(para, align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=None)
    run = para.add_run()
    _rpr_font(run, SZ_BODY)

    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar_begin)

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instrText)

    fldChar_sep = OxmlElement("w:fldChar")
    fldChar_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fldChar_sep)

    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar_end)

    hint = doc.add_paragraph()
    _pfmt(hint, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=None,
          space_before=Pt(4))
    r = hint.add_run("(Для обновления содержания нажмите Ctrl+A, затем F9)")
    _rpr_font(r, SZ_BODY, italic=True,
              color=RGBColor(0x80, 0x80, 0x80))

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Парсер Markdown → Document
# ---------------------------------------------------------------------------

def _is_table_line(line: str) -> bool:
    return "|" in line


def render_md(doc: Document, md_path: Path, *,
              skip_top_heading: bool = False) -> None:
    """Читает .md и добавляет форматированное содержимое в doc.

    Правила:
    - Таблица: накапливаем строки с |, потом ищем следующую строку
      «Таблица X.Y — ...» и вставляем её как подпись ПОСЛЕ таблицы.
    - Рисунок: строка ![alt](path) → вставляем изображение, следующая
      строка «Рисунок ...» → подпись.
    - Заголовки #/##/### → соответствующие стили.
    - Введение (## 1. Введение) — рендерим как раздел Введение.
    """
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    n = len(lines)
    i = 0

    in_code = False
    code_buf: list[str] = []

    # Очередь: таблица накоплена, ждём подпись
    pending_table: tuple[list[str], list[list[str]]] | None = None

    while i < n:
        line = lines[i]

        # ── Код ──────────────────────────────────────────────────────────
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                in_code = False
                add_code(doc, "\n".join(code_buf))
                code_buf = []
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # ── Подпись таблицы — ждём строку «Таблица X» ───────────────────
        if pending_table is not None:
            stripped = line.strip()
            if re.match(r"^Таблица\s+\d", stripped):
                add_table_with_caption(doc, pending_table[0], pending_table[1],
                                       stripped)
                pending_table = None
                i += 1
                continue
            elif stripped == "":
                i += 1
                continue
            else:
                # Подпись не найдена — выводим таблицу без подписи
                add_table_with_caption(doc, pending_table[0], pending_table[1], "")
                pending_table = None
                # не инкрементируем — обрабатываем текущую строку заново

        # ── Таблица ───────────────────────────────────────────────────────
        if _is_table_line(line):
            tbl_lines = [line]
            j = i + 1
            while j < n and _is_table_line(lines[j]):
                tbl_lines.append(lines[j])
                j += 1
            header, data = _parse_md_table(tbl_lines)
            if header:
                pending_table = (header, data)
            i = j
            continue

        # ── Рисунок ───────────────────────────────────────────────────────
        img_m = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line.strip())
        if img_m:
            alt_text = img_m.group(1)
            img_path = Path(img_m.group(2))
            # Следующая непустая строка — подпись «Рисунок ...»
            caption = alt_text
            j = i + 1
            while j < n and lines[j].strip() == "":
                j += 1
            if j < n and lines[j].strip().startswith("Рисунок"):
                caption = lines[j].strip()
                i = j + 1
            else:
                i += 1
            add_image_with_caption(doc, img_path, caption)
            continue

        stripped = line.strip()

        # ── Горизонтальная черта ──────────────────────────────────────────
        if re.match(r"^-{3,}$", stripped):
            i += 1
            continue

        # ── Заголовки ────────────────────────────────────────────────────
        if line.startswith("##### "):
            add_heading_sub4(doc, line[6:].strip())
        elif line.startswith("#### "):
            add_heading_sub4(doc, line[5:].strip())
        elif line.startswith("### "):
            add_heading_subsection(doc, line[4:].strip())
        elif line.startswith("## "):
            heading_text = line[3:].strip()
            # «1. Введение» в первой главе — отдельный блок (уже обработан)
            if re.match(r"^1\.\s*Введение", heading_text):
                pass  # пропускаем — Введение рендерится отдельным вызовом
            elif re.match(r"^Глава\s+\d", heading_text):
                pass  # пропускаем дублирующий заголовок главы
            else:
                add_heading_section(doc, heading_text)
        elif line.startswith("# "):
            if not skip_top_heading:
                pass  # мета-заголовок «ВКР магистра» — пропускаем

        # ── Маркированный список ─────────────────────────────────────────
        elif stripped.startswith("- ") or stripped.startswith("* "):
            add_bullet_item(doc, stripped[2:].strip())

        # ── Пустая строка ────────────────────────────────────────────────
        elif stripped == "":
            pass

        # ── Обычный текст ────────────────────────────────────────────────
        else:
            add_inline_para(doc, stripped)

        i += 1

    # Сбрасываем незакрытую таблицу
    if pending_table is not None:
        add_table_with_caption(doc, pending_table[0], pending_table[1], "")


# ---------------------------------------------------------------------------
# Введение — извлекаем из первой главы
# ---------------------------------------------------------------------------

def render_introduction(doc: Document, md_path: Path) -> None:
    """Вытаскивает раздел '## 1. Введение' и рендерит его отдельно."""
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    # Находим начало и конец Введения
    start = None
    end = None
    for idx, line in enumerate(lines):
        if re.match(r"^## 1\.\s*Введение", line.strip()):
            start = idx + 1
        elif start is not None and line.startswith("## "):
            end = idx
            break

    if start is None:
        return

    intro_lines = lines[start: end] if end else lines[start:]

    add_heading_chapter(doc, "ВВЕДЕНИЕ")

    n = len(intro_lines)
    i = 0
    pending_table: tuple[list[str], list[list[str]]] | None = None

    while i < n:
        line = intro_lines[i]

        if pending_table is not None:
            stripped = line.strip()
            if re.match(r"^Таблица\s+\d", stripped):
                add_table_with_caption(doc, pending_table[0], pending_table[1], stripped)
                pending_table = None
                i += 1
                continue
            elif stripped == "":
                i += 1
                continue
            else:
                add_table_with_caption(doc, pending_table[0], pending_table[1], "")
                pending_table = None

        if _is_table_line(line):
            tbl_lines = [line]
            j = i + 1
            while j < n and _is_table_line(intro_lines[j]):
                tbl_lines.append(intro_lines[j])
                j += 1
            header, data = _parse_md_table(tbl_lines)
            if header:
                pending_table = (header, data)
            i = j
            continue

        img_m = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line.strip())
        if img_m:
            caption = img_m.group(1)
            img_path = Path(img_m.group(2))
            j = i + 1
            while j < n and intro_lines[j].strip() == "":
                j += 1
            if j < n and intro_lines[j].strip().startswith("Рисунок"):
                caption = intro_lines[j].strip()
                i = j + 1
            else:
                i += 1
            add_image_with_caption(doc, img_path, caption)
            continue

        stripped = line.strip()
        if stripped.startswith("- ") or stripped.startswith("* "):
            add_bullet_item(doc, stripped[2:].strip())
        elif stripped == "":
            pass
        else:
            add_inline_para(doc, stripped)
        i += 1

    if pending_table is not None:
        add_table_with_caption(doc, pending_table[0], pending_table[1], "")


# ---------------------------------------------------------------------------
# Рендер «Глава X» — пропускает блок Введения
# ---------------------------------------------------------------------------

def render_chapter(doc: Document, md_path: Path) -> None:
    """Рендерит главу, пропуская мета-заголовки и блок Введения."""
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    # Пропускаем блок Введения (только в первой главе)
    skip_until = None
    for idx, line in enumerate(lines):
        if re.match(r"^## 1\.\s*Введение", line.strip()):
            # Найти конец введения
            for j in range(idx + 1, len(lines)):
                if lines[j].startswith("## ") and not re.match(r"^## 1\.\s*Введение", lines[j].strip()):
                    skip_until = j
                    break
            if skip_until is None:
                skip_until = len(lines)
            # Обрезаем: всё до ## 1. Введение + само Введение
            lines = lines[:idx] + lines[skip_until:]
            break

    # Записываем в временный объект и рендерим через основной парсер
    import tempfile, os
    with tempfile.NamedTemporaryFile(mode="w", encoding="utf-8",
                                     suffix=".md", delete=False) as tf:
        tf.write("\n".join(lines))
        tmp_path = Path(tf.name)
    try:
        render_md(doc, tmp_path, skip_top_heading=True)
    finally:
        os.unlink(tmp_path)


# ---------------------------------------------------------------------------
# Титульная страница
# ---------------------------------------------------------------------------

def add_title_page(doc: Document) -> None:
    def _c(text, size=14, bold=False, sa=0, sb=0):
        p = doc.add_paragraph()
        _pfmt(p, align=WD_ALIGN_PARAGRAPH.CENTER,
              first_indent=None, space_before=Pt(sb), space_after=Pt(sa))
        r = p.add_run(text)
        _rpr_font(r, Pt(size), bold=bold)

    _c("МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ", 12)
    _c("Федеральное государственное бюджетное образовательное учреждение", 12)
    _c("высшего образования", 12, sa=6)
    _c("РОССИЙСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ", 12, bold=True, sa=18)
    _c("Кафедра информационной безопасности", 13, sa=30)
    _c("ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА", 16, bold=True, sa=4)
    _c("(МАГИСТЕРСКАЯ ДИССЕРТАЦИЯ)", 14, bold=True, sa=20)
    _c("Тема:", 14, sa=4)
    _c("«Разработка системы детекции аномального сетевого трафика", 14, bold=True)
    _c("с использованием методов машинного обучения»", 14, bold=True, sa=30)

    fields = [
        ("Направление подготовки: ", "10.04.01 Информационная безопасность"),
        ("Профиль: ",                "Безопасность автоматизированных систем"),
        ("Обучающийся: ",            "_______________________"),
        ("Руководитель: ",           "_______________________"),
        ("Нормоконтроль: ",          "_______________________"),
    ]
    for label, value in fields:
        p = doc.add_paragraph()
        _pfmt(p, align=WD_ALIGN_PARAGRAPH.LEFT,
              first_indent=None, left_indent=Cm(8),
              space_before=Pt(0), space_after=Pt(3))
        r1 = p.add_run(label); _rpr_font(r1, Pt(13), bold=True)
        r2 = p.add_run(value); _rpr_font(r2, Pt(13))

    _c("", sa=20)
    _c("Москва — 2026", 13)
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Главная функция
# ---------------------------------------------------------------------------

def build_docx() -> None:
    doc = Document()
    set_page_margins(doc)
    add_page_numbers(doc)

    # Устанавливаем глобальный стиль Normal
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = SZ_BODY
    try:
        rPr = normal.element.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rFonts.set(qn(attr), FONT)
        existing = rPr.find(qn("w:rFonts"))
        if existing is not None:
            rPr.remove(existing)
        rPr.insert(0, rFonts)
    except Exception:
        pass

    # 1. Титульная страница
    add_title_page(doc)

    # 2. Содержание
    add_toc(doc)

    # 3. Введение (из первой главы)
    render_introduction(doc, CHAPTERS[0])
    doc.add_page_break()

    # 4. Глава 1
    add_heading_chapter(doc, "Глава 1. Аналитический и исследовательский раздел")
    render_chapter(doc, CHAPTERS[0])
    doc.add_page_break()

    # 5. Глава 2
    add_heading_chapter(doc, "Глава 2. Проектно-конструкторский раздел")
    render_md(doc, CHAPTERS[1])
    doc.add_page_break()

    # 6. Глава 3
    add_heading_chapter(doc,
        "Глава 3. Технологическая реализация и экспериментальная оценка системы")
    render_md(doc, CHAPTERS[2])
    doc.add_page_break()

    # 7. Список литературы
    add_heading_chapter(doc, "Список использованных источников")
    for ref_line in REFERENCES_TEXT.strip().splitlines():
        ref_line = ref_line.strip()
        if not ref_line:
            continue
        add_body(doc, ref_line, indent=True)

    doc.save(OUTPUT)
    print(f"Готово: {OUTPUT}  ({OUTPUT.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    build_docx()
