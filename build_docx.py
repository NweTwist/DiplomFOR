#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Конвертирует VKR_complete.md в DOCX по ГОСТ Р 7.0.11-2011.
Формирует автоматическое содержание через поле TOC.
"""
from __future__ import annotations
import re
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC    = Path("VKR_complete.md")
OUTPUT = Path("VKR_Dissertatsiya_final.docx")

FONT  = "Times New Roman"
SZ14  = Pt(14)
SZ16  = Pt(16)
SZ12  = Pt(12)
SZ10  = Pt(10)

LEFT  = Cm(3.0)
RIGHT = Cm(1.5)
TOP   = Cm(2.0)
BOT   = Cm(2.0)
INDENT = Cm(1.25)


# ─────────────────────────────── helpers ────────────────────────────────────

def _rpr(run, sz: Pt = SZ14, bold=False, italic=False,
         color: RGBColor | None = None):
    run.font.name  = FONT
    run.font.size  = sz
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    try:
        rPr = run._r.get_or_add_rPr()
        rf  = OxmlElement("w:rFonts")
        for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rf.set(qn(a), FONT)
        ex = rPr.find(qn("w:rFonts"))
        if ex is not None:
            rPr.remove(ex)
        rPr.insert(0, rf)
    except Exception:
        pass


def _pfmt(para, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
          first=INDENT, left=None,
          sb: Pt = Pt(0), sa: Pt = Pt(0),
          ls=WD_LINE_SPACING.ONE_POINT_FIVE):
    pf = para.paragraph_format
    pf.alignment         = align
    pf.space_before      = sb
    pf.space_after       = sa
    pf.line_spacing_rule = ls
    pf.first_line_indent = first
    if left is not None:
        pf.left_indent = left


def add_body(doc, text, *, ind=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             bold=False, italic=False, sz=SZ14):
    p = doc.add_paragraph()
    _pfmt(p, align=align, first=INDENT if ind else None)
    r = p.add_run(text)
    _rpr(r, sz=sz, bold=bold, italic=italic)


def add_bullet(doc, text):
    p = doc.add_paragraph()
    _pfmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
          first=Cm(-0.5), left=Cm(1.75))
    r = p.add_run(f"— {text}")
    _rpr(r)


def add_code(doc, text):
    p = doc.add_paragraph()
    _pfmt(p, align=WD_ALIGN_PARAGRAPH.LEFT, first=None,
          left=INDENT, sb=Pt(2), sa=Pt(2),
          ls=WD_LINE_SPACING.SINGLE)
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = SZ10


# ──────────────────── heading styles with Word outline levels ────────────────

def _set_outline_level(para, level: int):
    """Устанавливает outline level (0-based) через XML, чтобы TOC работало."""
    pPr = para._p.get_or_add_pPr()
    ol  = OxmlElement("w:outlineLvl")
    ol.set(qn("w:val"), str(level))
    # удаляем старый, если есть
    ex = pPr.find(qn("w:outlineLvl"))
    if ex is not None:
        pPr.remove(ex)
    pPr.append(ol)


def h1(doc, text):
    """Заголовок главы: 18pt, жирный, по ширине, отступ 1.25 см, 0 до / 12 после."""
    p = doc.add_paragraph()
    _pfmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first=INDENT,
          sb=Pt(0), sa=Pt(12))
    # Первая буква заглавная, остальные как есть (не all-caps)
    display = text[0].upper() + text[1:] if text else text
    r = p.add_run(display)
    _rpr(r, sz=Pt(18), bold=True)
    _set_outline_level(p, 0)


def h2(doc, text):          # X.Y Раздел — по левому, жирный, 14pt
    p = doc.add_paragraph()
    _pfmt(p, align=WD_ALIGN_PARAGRAPH.LEFT, first=None,
          sb=Pt(14), sa=Pt(6))
    r = p.add_run(text)
    _rpr(r, sz=SZ14, bold=True)
    _set_outline_level(p, 1)


def h3(doc, text):          # X.X.X Подраздел — жирный, без курсива
    p = doc.add_paragraph()
    _pfmt(p, align=WD_ALIGN_PARAGRAPH.LEFT, first=INDENT,
          sb=Pt(10), sa=Pt(4))
    r = p.add_run(text)
    _rpr(r, sz=SZ14, bold=True, italic=False)
    _set_outline_level(p, 2)


def h4(doc, text):          # подподраздел — жирный, без курсива
    p = doc.add_paragraph()
    _pfmt(p, align=WD_ALIGN_PARAGRAPH.LEFT, first=INDENT,
          sb=Pt(8), sa=Pt(2))
    r = p.add_run(text)
    _rpr(r, sz=SZ14, bold=True, italic=False)
    _set_outline_level(p, 3)


def add_image(doc, path_str, caption):
    ip = Path(path_str)
    p  = doc.add_paragraph()
    _pfmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, first=None,
          sb=Pt(6), sa=Pt(2))
    if ip.exists():
        try:
            p.add_run().add_picture(str(ip), width=Cm(15.0))
        except Exception:
            p.add_run(f"[Ошибка: {ip.name}]")
    else:
        p.add_run(f"[Файл не найден: ip]")
    cp = doc.add_paragraph()
    _pfmt(cp, align=WD_ALIGN_PARAGRAPH.CENTER, first=None,
          sb=Pt(0), sa=Pt(10))
    _rpr(cp.add_run(caption), sz=SZ12)


def parse_table(lines):
    rows = []
    for ln in lines:
        ln = ln.strip()
        if not ln:
            continue
        inner = ln.strip("|")
        if set(inner.replace(" ","").replace(":","")) <= set("-"):
            continue
        rows.append([c.strip() for c in ln.split("|") if c.strip()])
    return (rows[0], rows[1:]) if rows else ([], [])


def add_table(doc, header, data, caption):
    if not header:
        return
    # Название таблицы — ПЕРЕД таблицей
    if caption:
        cp = doc.add_paragraph()
        _pfmt(cp, align=WD_ALIGN_PARAGRAPH.LEFT, first=None,
              sb=Pt(6), sa=Pt(2))
        _rpr(cp.add_run(caption), sz=SZ12)
    t = doc.add_table(rows=1 + len(data), cols=len(header))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, h in enumerate(header):
        cell = t.rows[0].cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        pp = cell.paragraphs[0]
        _pfmt(pp, align=WD_ALIGN_PARAGRAPH.CENTER, first=None,
              sb=Pt(1), sa=Pt(1))
        _rpr(pp.add_run(h), sz=SZ12, bold=True)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "D9D9D9")
        cell._tc.get_or_add_tcPr().append(shd)
    for ri, row in enumerate(data):
        for ci, val in enumerate(row):
            if ci >= len(header): break
            cell = t.rows[ri+1].cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            pp = cell.paragraphs[0]
            _pfmt(pp, align=WD_ALIGN_PARAGRAPH.LEFT, first=None,
                  sb=Pt(1), sa=Pt(1))
            _rpr(pp.add_run(val), sz=SZ12)
    doc.add_paragraph()  # отступ после таблицы


def add_inline(doc, text, *, ind=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    _pfmt(p, align=align, first=INDENT if ind else None)
    segs = re.split(r"(`[^`\n]+`)", text)
    for s in segs:
        if s.startswith("`") and s.endswith("`"):
            r = p.add_run(s[1:-1])
            r.font.name = "Courier New"
            r.font.size = SZ12
        else:
            r = p.add_run(s)
            _rpr(r)


# ──────────────────── TOC ────────────────────────────────────────────────────

def add_toc(doc):
    h1(doc, "СОДЕРЖАНИЕ")
    p = doc.add_paragraph()
    _pfmt(p, align=WD_ALIGN_PARAGRAPH.LEFT, first=None)
    r = p.add_run()
    _rpr(r)
    for tag, text in [("begin", None), ("instrText", ' TOC \\o "1-3" \\h \\z \\u '), ("end", None)]:
        rr = p.add_run()
        _rpr(rr)
        if tag == "instrText":
            el = OxmlElement("w:instrText")
            el.set(qn("xml:space"), "preserve")
            el.text = text
            rr._r.append(el)
        else:
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), tag)
            rr._r.append(fc)
    hint = doc.add_paragraph()
    _pfmt(hint, align=WD_ALIGN_PARAGRAPH.CENTER, first=None, sb=Pt(4))
    _rpr(hint.add_run("(Ctrl+A, затем F9 для обновления содержания в Word)"),
         sz=SZ12, italic=True,
         color=RGBColor(0x80, 0x80, 0x80))
    doc.add_page_break()


# ──────────────────── нижний колонтитул с номером страницы ──────────────────

def add_page_numbers(doc):
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for tag, text in [("begin",None),("instrText","PAGE"),("end",None)]:
            rr = para.add_run()
            _rpr(rr)
            if tag == "instrText":
                el = OxmlElement("w:instrText")
                el.text = text
                rr._r.append(el)
            else:
                fc = OxmlElement("w:fldChar")
                fc.set(qn("w:fldCharType"), tag)
                rr._r.append(fc)


# ──────────────────── главный парсер MD → doc ────────────────────────────────

def render(doc, md_text: str):
    lines   = md_text.splitlines()
    n       = len(lines)
    i       = 0
    in_code = False
    code_buf: list[str] = []

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # ── код ──────────────────────────────────────────────────────────
        if stripped.startswith("```"):
            if not in_code:
                in_code  = True
                code_buf = []
            else:
                in_code = False
                add_code(doc, "\n".join(code_buf))
                code_buf = []
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # ── подпись таблицы ПЕРЕД таблицей ────────────────────────────────
        if re.match(r"^Таблица\s+", stripped) and "|" not in stripped:
            caption = stripped
            j = i + 1
            while j < n and lines[j].strip() == "":
                j += 1
            if j < n and lines[j].lstrip().startswith("|"):
                tbl_lines = []
                while j < n and lines[j].lstrip().startswith("|"):
                    tbl_lines.append(lines[j])
                    j += 1
                hdr, data = parse_table(tbl_lines)
                if hdr:
                    add_table(doc, hdr, data, caption)
                i = j
                continue
            else:
                add_inline(doc, stripped)
                i += 1
                continue

        # ── строка таблицы без подписи (на случай осиротевших) ───────────
        if stripped.startswith("|"):
            tbl_lines = []
            while i < n and lines[i].lstrip().startswith("|"):
                tbl_lines.append(lines[i])
                i += 1
            hdr, data = parse_table(tbl_lines)
            if hdr:
                add_table(doc, hdr, data, "")
            continue

        # ── рисунок ───────────────────────────────────────────────────────
        img = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', stripped)
        if img:
            caption = img.group(1)
            path    = img.group(2)
            j = i + 1
            while j < n and lines[j].strip() == "":
                j += 1
            if j < n and lines[j].strip().startswith("Рисунок"):
                caption = lines[j].strip()
                i = j + 1
            else:
                i += 1
            add_image(doc, path, caption)
            continue

        # ── разделитель ───────────────────────────────────────────────────
        if re.match(r"^-{3,}$", stripped):
            i += 1
            continue

        # ── заголовки ────────────────────────────────────────────────────
        if line.startswith("##### "):
            h4(doc, line[6:].strip())
        elif line.startswith("#### "):
            h4(doc, line[5:].strip())
        elif line.startswith("### "):
            h3(doc, line[4:].strip())
        elif line.startswith("## "):
            txt = line[3:].strip()
            if txt in ("Введение", "Заключение", "Список использованных источников") \
               or re.match(r"^Глава\s+\d", txt):
                h1(doc, txt)
            else:
                h2(doc, txt)
        elif line.startswith("# "):
            h1(doc, line[2:].strip())
        elif stripped.startswith("- ") or stripped.startswith("* "):
            add_bullet(doc, stripped[2:].strip())
        elif stripped == "":
            pass
        else:
            add_inline(doc, stripped)

        i += 1


# ──────────────────── титул ──────────────────────────────────────────────────

def title_page(doc):
    def c(text, sz=14, bold=False, sa=0):
        p = doc.add_paragraph()
        _pfmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, first=None,
              sa=Pt(sa))
        _rpr(p.add_run(text), sz=Pt(sz), bold=bold)

    c("МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ", 12)
    c("Федеральное государственное бюджетное образовательное учреждение", 12)
    c("высшего образования", 12, sa=6)
    c("РОССИЙСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ", 12, bold=True, sa=18)
    c("Кафедра информационной безопасности", 13, sa=30)
    c("ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА", 16, bold=True, sa=4)
    c("(МАГИСТЕРСКАЯ ДИССЕРТАЦИЯ)", 14, bold=True, sa=20)
    c("Тема:", 14, sa=4)
    c("«Разработка системы детекции аномального сетевого трафика", 14, bold=True)
    c("с использованием методов машинного обучения»", 14, bold=True, sa=30)
    for label, val in [
        ("Направление подготовки: ", "10.04.01 Информационная безопасность"),
        ("Профиль: ",                "Безопасность автоматизированных систем"),
        ("Обучающийся: ",            "_______________________"),
        ("Руководитель: ",           "_______________________"),
        ("Нормоконтроль: ",          "_______________________"),
    ]:
        p = doc.add_paragraph()
        _pfmt(p, align=WD_ALIGN_PARAGRAPH.LEFT, first=None,
              left=Cm(8), sa=Pt(3))
        r1 = p.add_run(label);  _rpr(r1, sz=Pt(13), bold=True)
        r2 = p.add_run(val);    _rpr(r2, sz=Pt(13))
    c("", sa=20)
    c("Москва — 2026", 13)
    doc.add_page_break()


# ──────────────────── main ───────────────────────────────────────────────────

def main():
    md = SRC.read_text(encoding="utf-8")

    doc = Document()
    for sec in doc.sections:
        sec.left_margin   = LEFT
        sec.right_margin  = RIGHT
        sec.top_margin    = TOP
        sec.bottom_margin = BOT

    # Глобальный стиль Normal
    norm = doc.styles["Normal"]
    norm.font.name = FONT
    norm.font.size = SZ14
    try:
        rPr = norm.element.get_or_add_rPr()
        rf  = OxmlElement("w:rFonts")
        for a in ("w:ascii","w:hAnsi","w:eastAsia","w:cs"):
            rf.set(qn(a), FONT)
        ex = rPr.find(qn("w:rFonts"))
        if ex is not None: rPr.remove(ex)
        rPr.insert(0, rf)
    except Exception:
        pass

    add_page_numbers(doc)
    title_page(doc)
    add_toc(doc)
    render(doc, md)

    doc.save(OUTPUT)
    print(f"Готово: {OUTPUT}  ({OUTPUT.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    main()
