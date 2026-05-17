from docx import Document
from pathlib import Path
d = Document("VKR_Magisterskaya_Dissertatsiya.docx")
print(f"Абзацев: {len(d.paragraphs)}")
print(f"Таблиц:  {len(d.tables)}")
print(f"Размер:  {Path('VKR_Magisterskaya_Dissertatsiya.docx').stat().st_size // 1024} KB")
# первые 5 заголовков (outline level 0)
from docx.oxml.ns import qn
titles = []
for p in d.paragraphs:
    ol = p._p.find(f".//{qn('w:pPr')}/{qn('w:outlineLvl')}")
    if ol is not None and ol.get(qn("w:val")) == "0":
        titles.append(p.text[:70])
print(f"\nЗаголовки уровня 1 ({len(titles)} шт.):")
for t in titles:
    print(f"  {t}")
