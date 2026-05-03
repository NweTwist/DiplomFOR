#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""Export chapter 1 from a Markdown thesis file into a DOCX formatted close to GOST.

- Input: Markdown file (default: Диплом.md)
- Output: DOCX (default: Первая глава.docx)
- Extracts content starting from the first H2 that begins with "## 2." (Chapter 1 section)

Notes / limitations:
- LaTeX math ($...$, $$...$$) is preserved as plain text.
- Markdown support is intentionally limited to the constructs used in this repo:
  headings (#/##/###), paragraphs, ordered/unordered lists, pipe tables, fenced code blocks.
"""

from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


@dataclass
class GostConfig:
    font_name: str = "Times New Roman"
    font_size_pt: int = 14
    line_spacing: float = 1.5
    first_line_indent_cm: float = 1.25
    margin_left_cm: float = 3.0
    margin_right_cm: float = 1.5
    margin_top_cm: float = 2.0
    margin_bottom_cm: float = 2.0


_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_OL_RE = re.compile(r"^(\d+)[\.)]\s+(.*)$")
_UL_RE = re.compile(r"^[-*]\s+(.*)$")


def _set_run_font(run, font_name: str, size_pt: int, bold: Optional[bool] = None, italic: Optional[bool] = None):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic

    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)


def _configure_document_gost(doc: Document, cfg: GostConfig) -> None:
    section = doc.sections[0]
    section.left_margin = Cm(cfg.margin_left_cm)
    section.right_margin = Cm(cfg.margin_right_cm)
    section.top_margin = Cm(cfg.margin_top_cm)
    section.bottom_margin = Cm(cfg.margin_bottom_cm)

    style = doc.styles["Normal"]
    style.font.name = cfg.font_name
    style.font.size = Pt(cfg.font_size_pt)

    # Ensure East Asia font too
    style_element = style._element
    rPr = style_element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), cfg.font_name)
    rFonts.set(qn("w:hAnsi"), cfg.font_name)
    rFonts.set(qn("w:cs"), cfg.font_name)
    rFonts.set(qn("w:eastAsia"), cfg.font_name)


def _configure_paragraph_gost(p, cfg: GostConfig, *, first_line_indent: Optional[float] = None, align: Optional[int] = None):
    pf = p.paragraph_format
    pf.line_spacing = cfg.line_spacing
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    indent = cfg.first_line_indent_cm if first_line_indent is None else first_line_indent
    if indent and indent > 0:
        pf.first_line_indent = Cm(indent)
    else:
        pf.first_line_indent = None

    p.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY


def _add_page_number_footer(doc: Document, cfg: GostConfig) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add PAGE field
    run = p.add_run()
    _set_run_font(run, cfg.font_name, cfg.font_size_pt)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def _extract_first_chapter(md_text: str) -> str:
    lines = md_text.splitlines()

    start_idx: Optional[int] = None
    for i, line in enumerate(lines):
        if line.startswith("## 2."):
            start_idx = i
            break

    if start_idx is None:
        # Fallback: return whole document
        return md_text

    end_idx = len(lines)
    for j in range(start_idx + 1, len(lines)):
        if lines[j].startswith("## ") and not lines[j].startswith("## 2."):
            end_idx = j
            break

    return "\n".join(lines[start_idx:end_idx]).strip() + "\n"


def _split_inline_code(text: str) -> List[Tuple[str, bool]]:
    """Split a line by inline code spans marked with backticks.

    Returns list of (chunk, is_code).
    """
    parts: List[Tuple[str, bool]] = []
    if "`" not in text:
        return [(text, False)]

    buf = ""
    in_code = False
    for ch in text:
        if ch == "`":
            if buf:
                parts.append((buf, in_code))
                buf = ""
            in_code = not in_code
            continue
        buf += ch
    if buf:
        parts.append((buf, in_code))
    return parts


def _add_paragraph_with_inline(doc: Document, text: str, cfg: GostConfig, *, indent_cm: Optional[float] = None):
    p = doc.add_paragraph(style="Normal")
    _configure_paragraph_gost(p, cfg, first_line_indent=indent_cm)

    for chunk, is_code in _split_inline_code(text):
        if not chunk:
            continue
        run = p.add_run(chunk)
        if is_code:
            _set_run_font(run, "Courier New", cfg.font_size_pt)
        else:
            _set_run_font(run, cfg.font_name, cfg.font_size_pt)


def _is_table_header_sep(line: str) -> bool:
    s = line.strip()
    if "|" not in s:
        return False
    s = s.strip("|")
    cols = [c.strip() for c in s.split("|")]
    return all(re.fullmatch(r":?-{3,}:?", c) is not None for c in cols)


def _parse_table_block(lines: List[str], start: int) -> Tuple[Optional[List[List[str]]], int]:
    """Parse a GitHub-style pipe table starting at index start.

    Returns (rows, next_index) where rows includes header as first row.
    """
    if start + 1 >= len(lines):
        return None, start

    header = lines[start].strip()
    sep = lines[start + 1].strip()

    if "|" not in header or not _is_table_header_sep(sep):
        return None, start

    rows: List[List[str]] = []

    def split_row(row_line: str) -> List[str]:
        s = row_line.strip().strip("|")
        return [c.strip() for c in s.split("|")]

    rows.append(split_row(header))

    i = start + 2
    while i < len(lines):
        ln = lines[i]
        if not ln.strip():
            break
        if "|" not in ln:
            break
        rows.append(split_row(ln))
        i += 1

    return rows, i


def _add_table(doc: Document, rows: List[List[str]], cfg: GostConfig) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)

    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"

    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            txt = row[c_idx] if c_idx < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 else WD_ALIGN_PARAGRAPH.JUSTIFY

            run = p.add_run(txt)
            _set_run_font(run, cfg.font_name, cfg.font_size_pt, bold=(r_idx == 0))


def md_to_docx(md_text: str, out_path: Path, cfg: GostConfig) -> None:
    doc = Document()
    _configure_document_gost(doc, cfg)
    _add_page_number_footer(doc, cfg)

    lines = md_text.splitlines()
    i = 0

    in_code_block = False
    code_fence = "`````"  # placeholder, not used
    code_lines: List[str] = []

    while i < len(lines):
        line = lines[i]

        # Fenced code blocks
        if line.strip().startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_fence = line.strip()
                code_lines = []
            else:
                # close
                in_code_block = False
                if code_lines:
                    p = doc.add_paragraph(style="Normal")
                    _configure_paragraph_gost(p, cfg, first_line_indent=0, align=WD_ALIGN_PARAGRAPH.LEFT)
                    for idx, cl in enumerate(code_lines):
                        run = p.add_run((cl + "\n") if idx < len(code_lines) - 1 else cl)
                        _set_run_font(run, "Courier New", 12)
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line.rstrip("\n"))
            i += 1
            continue

        # Tables
        table_rows, next_i = _parse_table_block(lines, i)
        if table_rows is not None:
            _add_table(doc, table_rows, cfg)
            i = next_i
            continue

        # Headings
        m = _HEADING_RE.match(line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            p = doc.add_paragraph(style="Normal")
            if level == 1:
                _configure_paragraph_gost(p, cfg, first_line_indent=0, align=WD_ALIGN_PARAGRAPH.CENTER)
                run = p.add_run(text)
                _set_run_font(run, cfg.font_name, cfg.font_size_pt, bold=True)
            elif level == 2:
                _configure_paragraph_gost(p, cfg, first_line_indent=0, align=WD_ALIGN_PARAGRAPH.CENTER)
                run = p.add_run(text)
                _set_run_font(run, cfg.font_name, cfg.font_size_pt, bold=True)
            else:
                _configure_paragraph_gost(p, cfg, first_line_indent=0, align=WD_ALIGN_PARAGRAPH.LEFT)
                run = p.add_run(text)
                _set_run_font(run, cfg.font_name, cfg.font_size_pt, bold=True)
            i += 1
            continue

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Lists
        m_ol = _OL_RE.match(line)
        if m_ol:
            number = m_ol.group(1)
            content = m_ol.group(2).strip()
            _add_paragraph_with_inline(doc, f"{number}) {content}", cfg, indent_cm=0)
            i += 1
            continue

        m_ul = _UL_RE.match(line)
        if m_ul:
            content = m_ul.group(1).strip()
            _add_paragraph_with_inline(doc, f"— {content}", cfg, indent_cm=0)
            i += 1
            continue

        # Paragraph (merge following non-empty non-structural lines)
        para_lines = [line.strip()]
        j = i + 1
        while j < len(lines):
            nxt = lines[j]
            if not nxt.strip():
                break
            if nxt.strip().startswith("```"):
                break
            if _HEADING_RE.match(nxt):
                break
            if _OL_RE.match(nxt) or _UL_RE.match(nxt):
                break
            # table start
            if "|" in nxt and j + 1 < len(lines) and _is_table_header_sep(lines[j + 1]):
                break
            para_lines.append(nxt.strip())
            j += 1

        text = " ".join(para_lines)
        _add_paragraph_with_inline(doc, text, cfg)
        i = j

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main() -> None:
    parser = argparse.ArgumentParser(description="Export first chapter from Markdown to DOCX (GOST-like formatting)")
    parser.add_argument("--md", default="Диплом.md", help="Input Markdown file")
    parser.add_argument("--out", default="Первая глава.docx", help="Output DOCX file")
    parser.add_argument(
        "--whole",
        action="store_true",
        help="Export the whole Markdown instead of extracting chapter starting from '## 2.'",
    )

    args = parser.parse_args()

    md_path = Path(args.md)
    out_path = Path(args.out)

    text = md_path.read_text(encoding="utf-8")
    if not args.whole:
        text = _extract_first_chapter(text)

    cfg = GostConfig()
    md_to_docx(text, out_path, cfg)
    print(f"Saved: {out_path.resolve()}")


if __name__ == "__main__":
    main()
